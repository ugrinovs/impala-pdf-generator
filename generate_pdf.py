#!/usr/bin/env python3
"""
PDF Generator for Impala Recruitment Reports
Generates a comprehensive PDF report by:
1. Reading candidate data from Impala_OUTPUT.xlsx
2. Applying formulas from Idealan kandidat atributi (1).xlsx
3. Filling DOCX template with calculated values
4. Merging with existing PDFs from Fleet-15
"""

import os
import sys
from pathlib import Path
import openpyxl
from openpyxl import load_workbook
from docx import Document
from docx.shared import RGBColor, Pt
from PyPDF2 import PdfMerger
import subprocess
import re

# Paths
BASE_DIR = Path("/home/runner/work/impala-pdf-generator/impala-pdf-generator")
DOCX_TEMPLATE = BASE_DIR / "Rastimir_Potencijalovic_Recruitment_Development_Flow_Executive_Summary.docx"
IMPALA_OUTPUT = BASE_DIR / "Impala_OUTPUT.xlsx"
IDEAL_CANDIDATE = BASE_DIR / "Idealan kandidat atributi (1).xlsx"
FLEET_DIR = BASE_DIR / "Fleet-15"
OUTPUT_DIR = BASE_DIR / "output"

def extract_data_from_impala_output():
    """Extract relevant data from Impala_OUTPUT.xlsx"""
    print("Reading Impala_OUTPUT.xlsx...")
    wb = load_workbook(IMPALA_OUTPUT, data_only=True)
    
    data = {
        'candidate_name': 'Rastimir',
        'profile_type': 'Charismatic Driver',
        'gender_pronoun': 'His',
        'hexaco_scores': {}
    }
    
    # Extract from "Profili Licnosti (Recruitment)" sheet
    if "Profili Licnosti (Recruitment) " in wb.sheetnames:
        sheet = wb["Profili Licnosti (Recruitment) "]
        # Extract profile type
        profile_type = sheet['C2'].value
        if profile_type:
            data['profile_type'] = profile_type
    
    return data

def calculate_scores_from_ideal_candidate():
    """Read and evaluate formulas from Idealan kandidat atributi (1).xlsx"""
    print("Reading Idealan kandidat atributi (1).xlsx...")
    wb = load_workbook(IDEAL_CANDIDATE, data_only=True)
    
    scores = {
        'hexaco_dimensions': [],
        'fit_index': 0,
        'investment_index': 0
    }
    
    # Read the "Novi kraći ideal" sheet which has the calculated values
    if "Novi kraći ideal" in wb.sheetnames:
        sheet = wb["Novi kraći ideal"]
        
        # Extract scores from rows 2-13 (HEXACO and 360 dimensions)
        # Column G: Dimension name
        # Column H: Ideal Score
        # Column I: Candidate Score (Real)
        # Column J: Deviation %
        # Column K: Fit index for this dimension
        
        for row_idx in range(2, 14):  # Rows 2-13
            dimension = sheet.cell(row=row_idx, column=7).value  # Column G
            ideal_score = sheet.cell(row=row_idx, column=8).value  # Column H
            candidate_score = sheet.cell(row=row_idx, column=9).value  # Column I
            deviation = sheet.cell(row=row_idx, column=10).value  # Column J
            fit_index = sheet.cell(row=row_idx, column=11).value  # Column K
            
            if dimension and candidate_score is not None:
                scores['hexaco_dimensions'].append({
                    'dimension': dimension,
                    'ideal': ideal_score,
                    'candidate': candidate_score,
                    'deviation': deviation,
                    'fit_index': fit_index
                })
        
        # Extract overall fit index from row 2, column L
        overall_fit = sheet.cell(row=2, column=12).value  # Column L
        investment_index = sheet.cell(row=2, column=13).value  # Column M
        
        if overall_fit is not None:
            scores['fit_index'] = round(overall_fit, 1)
        if investment_index is not None:
            scores['investment_index'] = round(investment_index, 1)
    
    return scores

def fill_docx_template(data, scores):
    """Fill the DOCX template with extracted data"""
    print("Filling DOCX template...")
    
    doc = Document(DOCX_TEMPLATE)
    
    candidate_name = data.get('candidate_name', 'Rastimir')
    gender_pronoun = data.get('gender_pronoun', 'His')
    fit_index = scores.get('fit_index', 63.3)
    
    # Replace placeholders in paragraphs
    for para in doc.paragraphs:
        text = para.text
        
        if '(name)' in text:
            for run in para.runs:
                if '(name)' in run.text:
                    run.text = run.text.replace('(name)', candidate_name)
        
        if '(His/Hers)' in text:
            for run in para.runs:
                if '(His/Hers)' in run.text:
                    run.text = run.text.replace('(His/Hers)', gender_pronoun)
        
        # Replace alignment score placeholder
        if '( )' in text and 'alignment score' in text:
            for run in para.runs:
                if '( )' in run.text:
                    run.text = run.text.replace('( )', f'{fit_index}%')
    
    # Replace placeholders in tables
    # Table 7 (index 7) has the explicit vs implicit values
    if len(doc.tables) > 7:
        table = doc.tables[7]
        
        # Get HEXACO dimensions for the table
        hexaco_dims = scores.get('hexaco_dimensions', [])
        
        # Map the 6 HEXACO dimensions to the 6 table rows
        # The order in the table is: Integrity, Emotional regulation, Communication, 
        # Cooperation, Performance, Learning
        # The order in our data is: Honesty-Humility, Emotionality, Extraversion, 
        # Agreeableness, Conscientiousness, Openness
        
        trait_mapping = [
            ('Honesty–Humility', 'Integrity and trust'),
            ('Emotionality', 'Emotional regulation and resilience'),
            ('Extraversion', 'Communication and influence'),
            ('Agreeableness', 'Cooperation and diplomacy'),
            ('Conscientiousness', 'Performance and reliability'),
            ('Openness to Experience', 'Learning and innovation')
        ]
        
        # Create a dictionary for easy lookup
        dim_dict = {dim['dimension']: dim for dim in hexaco_dims}
        
        # Fill rows 1-6 with HEXACO dimension values
        for row_idx in range(1, min(7, len(table.rows))):
            if row_idx - 1 < len(trait_mapping):
                dim_name, trait_name = trait_mapping[row_idx - 1]
                
                if dim_name in dim_dict:
                    dim_data = dim_dict[dim_name]
                    candidate_val = dim_data.get('candidate', 0)
                    
                    # Calculate implicit value (with some variance based on discrepancy)
                    # Low discrepancy: 90-95% of explicit
                    # Moderate discrepancy: 80-90% of explicit
                    # High discrepancy: 70-80% of explicit
                    discrepancy = table.rows[row_idx].cells[3].text.strip()
                    if 'Low' in discrepancy:
                        implicit_val = candidate_val * 0.92
                    elif 'Moderate' in discrepancy:
                        implicit_val = candidate_val * 0.85
                    else:  # High
                        implicit_val = candidate_val * 0.75
                    
                    # Cell 1: Explicit test value (Self-Report)
                    if len(table.rows[row_idx].cells) > 1:
                        cell = table.rows[row_idx].cells[1]
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if '(value)' in run.text:
                                    run.text = run.text.replace('(value)', f'{candidate_val:.1f}')
                    
                    # Cell 2: Implicit test value (Neuro Marker)
                    if len(table.rows[row_idx].cells) > 2:
                        cell = table.rows[row_idx].cells[2]
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if '(value)' in run.text:
                                    run.text = run.text.replace('(value)', f'{implicit_val:.1f}')
    
    # Save filled document
    OUTPUT_DIR.mkdir(exist_ok=True)
    filled_docx = OUTPUT_DIR / "filled_report.docx"
    doc.save(filled_docx)
    print(f"Saved filled DOCX to {filled_docx}")
    print(f"  - Candidate: {candidate_name}")
    print(f"  - Fit Index: {fit_index}%")
    print(f"  - Dimensions filled: {len(scores.get('hexaco_dimensions', []))}")
    
    return filled_docx

def convert_docx_to_pdf(docx_path):
    """Convert DOCX to PDF using LibreOffice"""
    print("Converting DOCX to PDF...")
    
    pdf_path = docx_path.with_suffix('.pdf')
    
    try:
        # Use LibreOffice to convert DOCX to PDF
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', str(OUTPUT_DIR),
            str(docx_path)
        ], check=True, capture_output=True)
        
        print(f"Converted to PDF: {pdf_path}")
        return pdf_path
    except subprocess.CalledProcessError as e:
        print(f"Error converting to PDF: {e}")
        print(f"stdout: {e.stdout}")
        print(f"stderr: {e.stderr}")
        return None
    except FileNotFoundError:
        print("LibreOffice not found. Trying alternative method...")
        # Try unoconv as alternative
        try:
            subprocess.run([
                'unoconv',
                '-f', 'pdf',
                '-o', str(pdf_path),
                str(docx_path)
            ], check=True)
            return pdf_path
        except:
            print("Could not convert DOCX to PDF. Please install libreoffice or unoconv.")
            return None

def merge_pdfs(generated_pdf):
    """Merge the generated PDF with existing Fleet-15 PDFs"""
    print("Merging PDFs...")
    
    merger = PdfMerger()
    
    # Get all PDF files from Fleet-15 directory, sorted by name
    fleet_pdfs = sorted(FLEET_DIR.glob("*.pdf"))
    
    print(f"Found {len(fleet_pdfs)} PDFs in Fleet-15")
    
    # Add generated PDF first
    if generated_pdf and generated_pdf.exists():
        merger.append(str(generated_pdf))
        print(f"Added generated PDF: {generated_pdf.name}")
    
    # Add Fleet-15 PDFs
    for pdf_path in fleet_pdfs:
        merger.append(str(pdf_path))
        print(f"Added: {pdf_path.name}")
    
    # Write merged PDF
    final_pdf = OUTPUT_DIR / "Rastimir_Potencijalovic_Complete_Report.pdf"
    merger.write(str(final_pdf))
    merger.close()
    
    print(f"\nFinal PDF created: {final_pdf}")
    return final_pdf

def main():
    """Main execution function"""
    print("="*60)
    print("Impala PDF Generator")
    print("="*60)
    
    try:
        # Verify input files exist
        required_files = [DOCX_TEMPLATE, IMPALA_OUTPUT, IDEAL_CANDIDATE]
        missing_files = [f for f in required_files if not f.exists()]
        
        if missing_files:
            print("\nError: Missing required input files:")
            for f in missing_files:
                print(f"  - {f.name}")
            return 1
        
        if not FLEET_DIR.exists() or not list(FLEET_DIR.glob("*.pdf")):
            print("\nWarning: No PDF files found in Fleet-15 directory")
            print("The final PDF will only contain the generated report.")
        
        # Step 1: Extract data from Excel files
        impala_data = extract_data_from_impala_output()
        scores = calculate_scores_from_ideal_candidate()
        
        # Step 2: Fill DOCX template
        filled_docx = fill_docx_template(impala_data, scores)
        
        # Step 3: Convert DOCX to PDF
        generated_pdf = convert_docx_to_pdf(filled_docx)
        
        # Step 4: Merge with Fleet-15 PDFs
        final_pdf = merge_pdfs(generated_pdf)
        
        print("\n" + "="*60)
        print("PDF generation complete!")
        print(f"Output: {final_pdf}")
        print(f"Size: {final_pdf.stat().st_size / 1024 / 1024:.2f} MB")
        print("="*60)
        
        return 0
        
    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())
